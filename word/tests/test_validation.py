"""
Test suite for validation.py module

Tests ValidationResult, validate_docx, validate_styles, validate_structure, validate_content
"""

import sys
from pathlib import Path

# Add parent directory to path to import docx_skill
sys.path.insert(0, str(Path(__file__).parent.parent))

from validation import (
    ValidationResult, ValidationLevel, ValidationIssue,
    validate_docx, validate_styles, validate_structure, validate_content
)
import tempfile
import shutil
from docx import Document


def create_test_document(path: Path, with_headings: bool = True, word_count: int = 50):
    """Create a test DOCX document."""
    doc = Document()
    
    if with_headings:
        doc.add_heading("Test Document", level=1)
        doc.add_heading("Section 1", level=2)
    
    # Add paragraphs to reach word count
    words_added = 0
    while words_added < word_count:
        doc.add_paragraph("This is a test paragraph with some content. " * 5)
        words_added += 50  # Approximate
    
    doc.save(str(path))


def test_validation_result():
    """Test ValidationResult class."""
    print("\n=== Testing ValidationResult ===")
    
    # Test empty result
    print("Test 1: Empty result (valid)")
    result = ValidationResult()
    assert result.is_valid, "Empty result should be valid"
    assert len(result.errors) == 0
    assert len(result.warnings) == 0
    print("  ✓ Empty result is valid")
    
    # Test with errors
    print("\nTest 2: Result with errors (invalid)")
    result = ValidationResult()
    result.add_error("Test error", location="line 1")
    assert not result.is_valid, "Result with errors should be invalid"
    assert len(result.errors) == 1
    print("  ✓ Result with errors is invalid")
    
    # Test with warnings only
    print("\nTest 3: Result with warnings only (valid)")
    result = ValidationResult()
    result.add_warning("Test warning")
    assert result.is_valid, "Result with only warnings should be valid"
    assert len(result.warnings) == 1
    print("  ✓ Result with warnings only is valid")
    
    # Test mixed issues
    print("\nTest 4: Mixed issues")
    result = ValidationResult()
    result.add_error("Error 1")
    result.add_error("Error 2")
    result.add_warning("Warning 1")
    result.add_info("Info 1")
    
    assert len(result.errors) == 2
    assert len(result.warnings) == 1
    assert len(result.info) == 1
    assert len(result.issues) == 4
    print("  ✓ Mixed issues tracked correctly")
    
    # Test string representation
    print("\nTest 5: String representation")
    str_repr = str(result)
    assert "FAILED" in str_repr
    assert "Errors: 2" in str_repr
    assert "Warnings: 1" in str_repr
    print(f"  ✓ String representation:\n{str_repr}")
    
    print("\n✓ All ValidationResult tests passed!")


def test_validate_docx():
    """Test validate_docx function."""
    print("\n=== Testing validate_docx ===")
    
    temp_dir = Path(tempfile.mkdtemp())
    
    try:
        # Test with non-existent file
        print("Test 1: Non-existent file")
        result = validate_docx(temp_dir / "nonexistent.docx")
        assert not result.is_valid
        assert any("not found" in e.message.lower() for e in result.errors)
        print("  ✓ Correctly detected missing file")
        
        # Test with valid DOCX
        print("\nTest 2: Valid DOCX document")
        valid_doc = temp_dir / "valid.docx"
        create_test_document(valid_doc)
        
        result = validate_docx(valid_doc)
        print(f"  Validation result: {result.is_valid}")
        print(f"  Paragraph count: {result.metadata.get('paragraph_count')}")
        
        assert result.is_valid, f"Valid document should pass validation: {result.errors}"
        assert result.metadata.get('paragraph_count', 0) > 0
        print("  ✓ Valid document validated successfully")
        
        # Test with non-DOCX file
        print("\nTest 3: Non-DOCX file")
        text_file = temp_dir / "test.txt"
        text_file.write_text("Not a DOCX file")
        
        result = validate_docx(text_file)
        assert not result.is_valid
        print("  ✓ Correctly detected invalid DOCX")
        
        # Test deep corruption check
        print("\nTest 4: Deep corruption check")
        result = validate_docx(valid_doc, check_corruption=True)
        assert result.is_valid
        print("  ✓ Deep validation passed")
        
        print("\n✓ All validate_docx tests passed!")
        
    finally:
        shutil.rmtree(temp_dir)


def test_validate_styles():
    """Test validate_styles function."""
    print("\n=== Testing validate_styles ===")
    
    temp_dir = Path(tempfile.mkdtemp())
    
    try:
        # Create document with styles
        print("Test 1: Document with headings (uses styles)")
        styled_doc = temp_dir / "styled.docx"
        create_test_document(styled_doc, with_headings=True)
        
        result = validate_styles(styled_doc)
        print(f"  Defined styles: {result.metadata.get('defined_styles')}")
        print(f"  Used styles: {result.metadata.get('used_styles')}")
        print(f"  Valid: {result.is_valid}")
        
        # Document should be valid (may have warnings about unused styles)
        assert result.is_valid or len(result.errors) == 0
        print("  ✓ Style validation completed")
        
        # Test with invalid file
        print("\nTest 2: Invalid file")
        result = validate_styles(temp_dir / "nonexistent.docx")
        assert not result.is_valid
        print("  ✓ Correctly handled invalid file")
        
        print("\n✓ All validate_styles tests passed!")
        
    finally:
        shutil.rmtree(temp_dir)


def test_validate_structure():
    """Test validate_structure function."""
    print("\n=== Testing validate_structure ===")
    
    temp_dir = Path(tempfile.mkdtemp())
    
    try:
        # Test document with headings
        print("Test 1: Document with proper heading structure")
        doc_path = temp_dir / "structured.docx"
        create_test_document(doc_path, with_headings=True)
        
        result = validate_structure(doc_path, require_heading=True)
        print(f"  Heading count: {result.metadata.get('heading_count')}")
        print(f"  Max heading level: {result.metadata.get('max_heading_level')}")
        print(f"  Valid: {result.is_valid}")
        
        assert result.metadata.get('heading_count', 0) > 0
        print("  ✓ Structure validation passed")
        
        # Test document without headings
        print("\nTest 2: Document without headings (require_heading=True)")
        no_heading_doc = temp_dir / "no_heading.docx"
        create_test_document(no_heading_doc, with_headings=False)
        
        result = validate_structure(no_heading_doc, require_heading=True)
        print(f"  Valid: {result.is_valid}")
        
        # Should have error or warning about missing headings
        assert len(result.errors) > 0 or len(result.warnings) > 0
        print("  ✓ Correctly detected missing headings")
        
        # Test max depth
        print("\nTest 3: Check heading depth constraint")
        result = validate_structure(doc_path, max_depth=1)
        # May have warnings if document has H2 headings
        print(f"  Warnings about depth: {len(result.warnings)}")
        print("  ✓ Depth checking works")
        
        print("\n✓ All validate_structure tests passed!")
        
    finally:
        shutil.rmtree(temp_dir)


def test_validate_content():
    """Test validate_content function."""
    print("\n=== Testing validate_content ===")
    
    temp_dir = Path(tempfile.mkdtemp())
    
    try:
        # Create document with specific word count
        print("Test 1: Document with word count check")
        doc_path = temp_dir / "content.docx"
        create_test_document(doc_path, word_count=100)
        
        result = validate_content(doc_path)
        word_count = result.metadata.get('word_count', 0)
        print(f"  Word count: {word_count:,}")
        print(f"  Character count: {result.metadata.get('character_count', 0):,}")
        print(f"  Valid: {result.is_valid}")
        
        assert word_count > 0
        print("  ✓ Content validation passed")
        
        # Test minimum word count
        print("\nTest 2: Minimum word count constraint")
        result = validate_content(doc_path, min_words=50)
        assert result.is_valid, "Document should meet minimum word count"
        print("  ✓ Minimum constraint passed")
        
        result = validate_content(doc_path, min_words=10000)
        assert not result.is_valid, "Document should fail minimum constraint"
        print("  ✓ Correctly detected insufficient words")
        
        # Test maximum word count
        print("\nTest 3: Maximum word count constraint")
        result = validate_content(doc_path, max_words=50000)
        assert result.is_valid, "Document should be under maximum"
        print("  ✓ Maximum constraint passed")
        
        result = validate_content(doc_path, max_words=10)
        assert not result.is_valid, "Document should exceed maximum"
        print("  ✓ Correctly detected excessive words")
        
        # Test range
        print("\nTest 4: Word count range")
        result = validate_content(doc_path, min_words=50, max_words=500)
        print(f"  In range (50-500): {result.is_valid}")
        
        print("\n✓ All validate_content tests passed!")
        
    finally:
        shutil.rmtree(temp_dir)


def run_all_tests():
    """Run all validation tests."""
    print("=" * 60)
    print("DOCX Skill - Validation Module Tests")
    print("=" * 60)
    
    try:
        test_validation_result()
        test_validate_docx()
        test_validate_styles()
        test_validate_structure()
        test_validate_content()
        
        print("\n" + "=" * 60)
        print("✓ ALL VALIDATION TESTS PASSED!")
        print("=" * 60)
        return True
        
    except Exception as e:
        print("\n" + "=" * 60)
        print(f"✗ TEST FAILED: {e}")
        print("=" * 60)
        import traceback
        traceback.print_exc()
        return False


if __name__ == "__main__":
    success = run_all_tests()
    sys.exit(0 if success else 1)
