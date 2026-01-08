"""
Module: Advanced API - Full Document Control

Advanced document builder with complete control over styles, sections, tables,
and images. For users who need fine-grained control beyond the Simple API.

Public Interface:
    - AdvancedDocument: Advanced document builder with manager properties
    - StyleManager: Create and manage custom styles
    - SectionManager: Control sections, page layout, headers/footers
    - TableBuilder: Advanced table construction
    - ImageManager: Advanced image handling

Example:
    >>> from docx_skill.advanced import AdvancedDocument
    >>> 
    >>> # Create document with custom styles
    >>> doc = AdvancedDocument()
    >>> doc.styles.add_paragraph_style("Emphasis", font_size=14, italic=True)
    >>> doc.add_paragraph("Important text", style="Emphasis")
    >>> 
    >>> # Control sections and page layout
    >>> doc.sections.add_section(orientation="landscape")
    >>> doc.sections.set_margins(top=1.0, bottom=1.0, left=1.5, right=1.5)
    >>> 
    >>> # Advanced table features
    >>> table = doc.tables.create_table(rows=3, cols=4)
    >>> # Customize table...
    >>> 
    >>> doc.save("advanced.docx")
"""

from pathlib import Path
from typing import Optional, List, Any, Dict
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    from .ooxml import OOXMLDocument
    from .safety import SafeFileOperations
    from .validation import validate_docx
except ImportError:
    # Fallback for direct imports (e.g., in tests)
    from ooxml import OOXMLDocument
    from safety import SafeFileOperations
    from validation import validate_docx


class StyleManager:
    """Manage custom styles.
    
    Provides methods for creating and modifying paragraph and character styles.
    Allows full control over typography and formatting.
    
    Example:
        >>> doc = AdvancedDocument()
        >>> 
        >>> # Create custom paragraph style
        >>> doc.styles.add_paragraph_style(
        ...     name="CodeBlock",
        ...     font_name="Courier New",
        ...     font_size=10,
        ...     bold=False
        ... )
        >>> 
        >>> # Use custom style
        >>> doc.add_paragraph("print('Hello')", style="CodeBlock")
    """
    
    def __init__(self, document: Any):
        """Initialize with document.
        
        Args:
            document: python-docx Document object
        """
        self._document = document
    
    def add_paragraph_style(self, name: str, 
                           font_name: Optional[str] = None,
                           font_size: Optional[int] = None,
                           bold: bool = False,
                           italic: bool = False,
                           color: Optional[tuple] = None) -> Any:
        """Create custom paragraph style.
        
        Args:
            name: Style name (e.g., "MyCustomStyle")
            font_name: Font family (e.g., "Arial", "Times New Roman")
            font_size: Font size in points (e.g., 12)
            bold: Make text bold
            italic: Make text italic
            color: RGB color tuple (e.g., (255, 0, 0) for red)
        
        Returns:
            Created style object
            
        Example:
            >>> doc = AdvancedDocument()
            >>> # Create highlighted style
            >>> doc.styles.add_paragraph_style(
            ...     name="Highlight",
            ...     font_size=14,
            ...     bold=True,
            ...     color=(255, 0, 0)
            ... )
            >>> doc.add_paragraph("Important!", style="Highlight")
        """
        styles = self._document.styles
        
        # Check if style already exists
        try:
            existing_style = styles[name]
            # Style exists, modify it instead of adding
            style = existing_style
        except KeyError:
            # Style doesn't exist, create it
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        
        # Set font properties
        font = style.font
        if font_name:
            font.name = font_name
        if font_size:
            font.size = Pt(font_size)
        if bold:
            font.bold = True
        if italic:
            font.italic = True
        if color:
            font.color.rgb = RGBColor(*color)
        
        return style
    
    def add_character_style(self, name: str,
                          font_name: Optional[str] = None,
                          font_size: Optional[int] = None,
                          bold: bool = False,
                          italic: bool = False,
                          color: Optional[tuple] = None) -> Any:
        """Create custom character style.
        
        Character styles apply to individual runs of text within paragraphs,
        unlike paragraph styles which apply to entire paragraphs.
        
        Args:
            name: Style name
            font_name: Font family
            font_size: Font size in points
            bold: Make text bold
            italic: Make text italic
            color: RGB color tuple
        
        Returns:
            Created style object
            
        Example:
            >>> doc = AdvancedDocument()
            >>> # Create code style for inline code
            >>> doc.styles.add_character_style(
            ...     name="InlineCode",
            ...     font_name="Courier New",
            ...     font_size=11,
            ...     color=(0, 0, 128)
            ... )
        """
        styles = self._document.styles
        style = styles.add_style(name, WD_STYLE_TYPE.CHARACTER)
        
        # Set font properties
        font = style.font
        if font_name:
            font.name = font_name
        if font_size:
            font.size = Pt(font_size)
        if bold:
            font.bold = True
        if italic:
            font.italic = True
        if color:
            font.color.rgb = RGBColor(*color)
        
        return style
    
    def list_styles(self) -> List[str]:
        """List all available styles.
        
        Returns:
            List of style names
            
        Example:
            >>> doc = AdvancedDocument()
            >>> styles = doc.styles.list_styles()
            >>> print(f"Available styles: {', '.join(styles[:5])}")
        """
        return [style.name for style in self._document.styles]
    
    def get_style(self, name: str) -> Any:
        """Get style by name.
        
        Args:
            name: Style name
        
        Returns:
            Style object or None if not found
            
        Example:
            >>> doc = AdvancedDocument()
            >>> heading_style = doc.styles.get_style("Heading 1")
            >>> if heading_style:
            ...     print(f"Font: {heading_style.font.name}")
        """
        try:
            return self._document.styles[name]
        except KeyError:
            return None


class SectionManager:
    """Manage sections and page setup.
    
    Sections control page layout, orientation, margins, headers, and footers.
    Each section can have different page settings.
    
    Example:
        >>> doc = AdvancedDocument()
        >>> 
        >>> # Set margins for current section
        >>> doc.sections.set_margins(top=1.0, bottom=1.0, left=1.5, right=1.5)
        >>> 
        >>> # Add landscape section for wide table
        >>> doc.sections.add_section(orientation="landscape")
        >>> doc.add_paragraph("This is in landscape mode")
        >>> 
        >>> # Add header/footer
        >>> doc.sections.add_header("Confidential Document")
        >>> doc.sections.add_footer("Page")
    """
    
    def __init__(self, document: Any):
        """Initialize with document.
        
        Args:
            document: python-docx Document object
        """
        self._document = document
    
    def add_section(self, 
                   orientation: str = "portrait",
                   page_width: Optional[float] = None,
                   page_height: Optional[float] = None) -> Any:
        """Add new section with layout.
        
        Args:
            orientation: 'portrait' or 'landscape'
            page_width: Width in inches (None = default Letter: 8.5")
            page_height: Height in inches (None = default Letter: 11")
        
        Returns:
            New Section object
            
        Example:
            >>> doc = AdvancedDocument()
            >>> doc.add_paragraph("Portrait page")
            >>> 
            >>> # Add landscape section
            >>> doc.sections.add_section(orientation="landscape")
            >>> doc.add_paragraph("Landscape page")
            >>> 
            >>> # Custom page size (A4)
            >>> doc.sections.add_section(
            ...     page_width=8.27,
            ...     page_height=11.69
            ... )
        """
        # Add section break
        new_section = self._document.add_section()
        
        # Set orientation
        if orientation.lower() == "landscape":
            # Swap width and height for landscape
            new_section.orientation = 1  # WD_ORIENT.LANDSCAPE
            new_section.page_width = Inches(page_height or 11)
            new_section.page_height = Inches(page_width or 8.5)
        else:
            # Portrait (default)
            new_section.orientation = 0  # WD_ORIENT.PORTRAIT
            new_section.page_width = Inches(page_width or 8.5)
            new_section.page_height = Inches(page_height or 11)
        
        return new_section
    
    def set_margins(self, top: float, bottom: float,
                   left: float, right: float) -> None:
        """Set page margins in inches.
        
        Sets margins for the current (last) section.
        
        Args:
            top: Top margin in inches
            bottom: Bottom margin in inches
            left: Left margin in inches
            right: Right margin in inches
        
        Example:
            >>> doc = AdvancedDocument()
            >>> # Narrow margins
            >>> doc.sections.set_margins(
            ...     top=0.5, bottom=0.5,
            ...     left=0.5, right=0.5
            ... )
            >>> 
            >>> # Wide left margin for binding
            >>> doc.sections.set_margins(
            ...     top=1.0, bottom=1.0,
            ...     left=2.0, right=1.0
            ... )
        """
        section = self._document.sections[-1]
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)
    
    def add_header(self, text: str) -> None:
        """Add header to current section.
        
        Args:
            text: Header text
        
        Example:
            >>> doc = AdvancedDocument()
            >>> doc.sections.add_header("Company Confidential")
            >>> doc.add_paragraph("Document content...")
        """
        section = self._document.sections[-1]
        header = section.header
        header.paragraphs[0].text = text
    
    def add_footer(self, text: str) -> None:
        """Add footer to current section.
        
        Args:
            text: Footer text
        
        Example:
            >>> doc = AdvancedDocument()
            >>> doc.sections.add_footer("© 2024 Company Name")
            >>> doc.add_paragraph("Document content...")
        """
        section = self._document.sections[-1]
        footer = section.footer
        footer.paragraphs[0].text = text


class TableBuilder:
    """Advanced table construction.
    
    Provides fine-grained control over table structure, styling, and content.
    
    Example:
        >>> doc = AdvancedDocument()
        >>> 
        >>> # Create empty table and customize
        >>> table = doc.tables.create_table(rows=3, cols=4)
        >>> table.rows[0].cells[0].text = "Header 1"
        >>> 
        >>> # Create table from data
        >>> table = doc.tables.add_table_from_data(
        ...     data=[["A", "B"], ["C", "D"]],
        ...     headers=["Col 1", "Col 2"],
        ...     style="Light Grid Accent 1"
        ... )
    """
    
    def __init__(self, document: Any):
        """Initialize with document.
        
        Args:
            document: python-docx Document object
        """
        self._document = document
    
    def create_table(self, rows: int, cols: int) -> Any:
        """Create table structure.
        
        Args:
            rows: Number of rows
            cols: Number of columns
        
        Returns:
            python-docx Table object
            
        Example:
            >>> doc = AdvancedDocument()
            >>> table = doc.tables.create_table(rows=5, cols=3)
            >>> 
            >>> # Populate cells
            >>> for i in range(5):
            ...     for j in range(3):
            ...         table.rows[i].cells[j].text = f"R{i}C{j}"
        """
        return self._document.add_table(rows=rows, cols=cols)
    
    def add_table_from_data(self, data: List[List[str]], 
                           headers: Optional[List[str]] = None,
                           style: Optional[str] = None) -> Any:
        """Create table from 2D array with optional headers and style.
        
        Args:
            data: 2D list of cell values
            headers: Optional header row
            style: Table style name (e.g., "Light Grid Accent 1")
        
        Returns:
            python-docx Table object
            
        Example:
            >>> doc = AdvancedDocument()
            >>> 
            >>> # Sales data table
            >>> table = doc.tables.add_table_from_data(
            ...     data=[
            ...         ["Q1", "100", "150"],
            ...         ["Q2", "120", "180"],
            ...         ["Q3", "140", "200"]
            ...     ],
            ...     headers=["Quarter", "Sales", "Target"],
            ...     style="Medium Grid 1 Accent 1"
            ... )
        """
        if not data:
            return None
        
        # Determine dimensions
        num_cols = len(data[0])
        num_rows = len(data) + (1 if headers else 0)
        
        # Create table
        table = self._document.add_table(rows=num_rows, cols=num_cols)
        
        # Apply style if specified
        if style:
            try:
                table.style = style
            except KeyError:
                # Style doesn't exist, skip
                pass
        
        # Add headers if provided
        row_offset = 0
        if headers:
            header_row = table.rows[0]
            for col_idx, header_text in enumerate(headers):
                cell = header_row.cells[col_idx]
                cell.text = header_text
                # Make header bold
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            row_offset = 1
        
        # Add data rows
        for row_idx, row_data in enumerate(data):
            table_row = table.rows[row_idx + row_offset]
            for col_idx, cell_value in enumerate(row_data):
                table_row.cells[col_idx].text = str(cell_value)
        
        return table
    
    def list_table_styles(self) -> List[str]:
        """List available table styles.
        
        Returns:
            List of table style names
            
        Example:
            >>> doc = AdvancedDocument()
            >>> styles = doc.tables.list_table_styles()
            >>> print(f"Found {len(styles)} table styles")
            >>> print(styles[:5])  # Show first 5
        """
        return [style.name for style in self._document.styles 
                if style.type == WD_STYLE_TYPE.TABLE]


class ImageManager:
    """Advanced image handling.
    
    Provides control over image sizing, positioning, and properties.
    
    Example:
        >>> doc = AdvancedDocument()
        >>> 
        >>> # Add image with specific width
        >>> doc.images.add_image("photo.jpg", width_inches=4.0)
        >>> 
        >>> # Add image with width and height
        >>> doc.images.add_image(
        ...     "chart.png",
        ...     width_inches=5.0,
        ...     height_inches=3.0
        ... )
    """
    
    def __init__(self, document: Any):
        """Initialize with document.
        
        Args:
            document: python-docx Document object
        """
        self._document = document
    
    def add_image(self, path: str,
                 width_inches: Optional[float] = None,
                 height_inches: Optional[float] = None) -> Any:
        """Add image with sizing.
        
        Args:
            path: Path to image file
            width_inches: Width in inches (None = original)
            height_inches: Height in inches (None = proportional to width)
        
        Returns:
            InlineShape object
            
        Raises:
            FileNotFoundError: If image file doesn't exist
            
        Example:
            >>> doc = AdvancedDocument()
            >>> 
            >>> # Original size
            >>> doc.images.add_image("logo.png")
            >>> 
            >>> # Specific width (proportional height)
            >>> doc.images.add_image("banner.jpg", width_inches=6.0)
            >>> 
            >>> # Specific width and height
            >>> doc.images.add_image(
            ...     "photo.jpg",
            ...     width_inches=4.0,
            ...     height_inches=3.0
            ... )
        """
        image_path = Path(path)
        if not image_path.exists():
            raise FileNotFoundError(f"Image not found: {path}")
        
        # Add image with sizing
        if width_inches and height_inches:
            return self._document.add_picture(
                str(image_path),
                width=Inches(width_inches),
                height=Inches(height_inches)
            )
        elif width_inches:
            return self._document.add_picture(
                str(image_path),
                width=Inches(width_inches)
            )
        elif height_inches:
            return self._document.add_picture(
                str(image_path),
                height=Inches(height_inches)
            )
        else:
            return self._document.add_picture(str(image_path))


class AdvancedDocument:
    """Advanced document builder with full styling control.
    
    Provides complete control over document structure, styles, sections,
    tables, and images through specialized manager classes.
    
    Use this API when you need fine-grained control beyond the Simple API.
    For basic document creation, use DocumentBuilder instead.
    
    Attributes:
        styles: StyleManager for custom styles
        sections: SectionManager for page layout
        tables: TableBuilder for advanced tables
        images: ImageManager for image handling
    
    Example:
        >>> doc = AdvancedDocument()
        >>> 
        >>> # Custom styling
        >>> doc.styles.add_paragraph_style(
        ...     "Emphasis",
        ...     font_size=14,
        ...     bold=True,
        ...     color=(255, 0, 0)
        ... )
        >>> 
        >>> # Page layout
        >>> doc.sections.set_margins(top=1.0, bottom=1.0, left=1.5, right=1.5)
        >>> doc.sections.add_header("My Document")
        >>> 
        >>> # Content
        >>> doc.add_heading("Introduction", level=1)
        >>> doc.add_paragraph("Important point", style="Emphasis")
        >>> 
        >>> # Advanced table
        >>> table = doc.tables.add_table_from_data(
        ...     data=[["A", "B"], ["C", "D"]],
        ...     headers=["X", "Y"],
        ...     style="Light Grid"
        ... )
        >>> 
        >>> # Save
        >>> doc.save("advanced.docx")
    """
    
    def __init__(self, template: Optional[str] = None):
        """Create advanced document.
        
        Args:
            template: Optional template name or path. If None, uses modern
                     Microsoft 365 template with Aptos font.
        """
        if template is None or template == "modern":
            self._ooxml_doc = OOXMLDocument(use_modern_template=True)
        elif template == "legacy":
            self._ooxml_doc = OOXMLDocument(use_modern_template=False)
        else:
            self._ooxml_doc = OOXMLDocument.load(template)
        
        # Initialize managers
        self._styles = StyleManager(self._ooxml_doc.document)
        self._sections = SectionManager(self._ooxml_doc.document)
        self._tables = TableBuilder(self._ooxml_doc.document)
        self._images = ImageManager(self._ooxml_doc.document)
        
        self._safe_ops = SafeFileOperations(default_allow_overwrite=False)
    
    @property
    def styles(self) -> StyleManager:
        """Access style management.
        
        Returns:
            StyleManager instance
            
        Example:
            >>> doc = AdvancedDocument()
            >>> doc.styles.add_paragraph_style("Custom", font_size=16)
        """
        return self._styles
    
    @property
    def sections(self) -> SectionManager:
        """Access section management.
        
        Returns:
            SectionManager instance
            
        Example:
            >>> doc = AdvancedDocument()
            >>> doc.sections.set_margins(1.0, 1.0, 1.0, 1.0)
        """
        return self._sections
    
    @property
    def tables(self) -> TableBuilder:
        """Access advanced table features.
        
        Returns:
            TableBuilder instance
            
        Example:
            >>> doc = AdvancedDocument()
            >>> table = doc.tables.create_table(rows=3, cols=4)
        """
        return self._tables
    
    @property
    def images(self) -> ImageManager:
        """Access advanced image features.
        
        Returns:
            ImageManager instance
            
        Example:
            >>> doc = AdvancedDocument()
            >>> doc.images.add_image("photo.jpg", width_inches=5.0)
        """
        return self._images
    
    def add_paragraph(self, text: str, style: Optional[str] = None) -> Any:
        """Add paragraph with optional style.
        
        Args:
            text: Paragraph text
            style: Style name (built-in or custom)
        
        Returns:
            python-docx Paragraph object
            
        Example:
            >>> doc = AdvancedDocument()
            >>> doc.add_paragraph("Normal text")
            >>> doc.add_paragraph("Heading text", style="Heading 1")
            >>> doc.add_paragraph("Custom styled", style="MyCustomStyle")
        """
        return self._ooxml_doc.add_paragraph(text, style)
    
    def add_heading(self, text: str, level: int = 1) -> Any:
        """Add heading.
        
        Args:
            text: Heading text
            level: Heading level (1-9)
        
        Returns:
            python-docx Paragraph object
            
        Example:
            >>> doc = AdvancedDocument()
            >>> doc.add_heading("Chapter 1", level=1)
            >>> doc.add_heading("Section 1.1", level=2)
        """
        return self._ooxml_doc.add_heading(text, level)
    
    def add_page_break(self) -> Any:
        """Add page break.
        
        Returns:
            python-docx Paragraph object containing page break
            
        Example:
            >>> doc = AdvancedDocument()
            >>> doc.add_paragraph("Page 1")
            >>> doc.add_page_break()
            >>> doc.add_paragraph("Page 2")
        """
        return self._ooxml_doc.add_page_break()
    
    def save(self, output_path: str, overwrite: bool = False) -> str:
        """Save with validation.
        
        Args:
            output_path: Where to save document
            overwrite: If False, raises error if file exists
        
        Returns:
            Path where document was saved
            
        Raises:
            FileExistsError: If file exists and overwrite=False
            
        Example:
            >>> doc = AdvancedDocument()
            >>> doc.add_heading("Document")
            >>> path = doc.save("output.docx")
        """
        # Save to temp location first
        temp_path = Path(output_path).parent / f"~temp_{Path(output_path).name}"
        self._ooxml_doc.save(temp_path)
        
        # Validate the saved document
        try:
            validation = validate_docx(str(temp_path))
            if not validation.is_valid:
                for issue in validation.issues:
                    print(f"Warning: {issue.message}")
        except Exception as e:
            # Don't fail save on validation errors
            print(f"Warning: Validation failed: {e}")
        
        # Use SafeFileOperations to move to final location
        try:
            data = temp_path.read_bytes()
            self._safe_ops.write_file(
                data=data,
                target_path=output_path,
                allow_overwrite=overwrite,
                backup=overwrite
            )
        finally:
            if temp_path.exists():
                temp_path.unlink()
        
        return str(output_path)
    
    def get_ooxml_document(self) -> OOXMLDocument:
        """Get underlying OOXMLDocument for raw OOXML operations.
        
        Returns:
            OOXMLDocument instance
            
        Example:
            >>> doc = AdvancedDocument()
            >>> ooxml = doc.get_ooxml_document()
            >>> # Direct OOXML manipulation
            >>> body = ooxml.get_body_element()
        """
        return self._ooxml_doc
    
    def __repr__(self) -> str:
        """String representation."""
        para_count = len(self._ooxml_doc.get_paragraphs())
        table_count = len(self._ooxml_doc.get_tables())
        return f"AdvancedDocument({para_count} paragraphs, {table_count} tables)"
