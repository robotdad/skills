"""
Module: OOXML Layer

Low-level OOXML (Office Open XML) access layer for direct document manipulation.
Wraps python-docx Document class and provides direct access to underlying XML
structure for advanced operations.

This is the foundation layer that higher-level APIs (simple.py, advanced.py)
build upon. Use this layer when you need precise control over document structure
or when working with features not exposed by higher-level APIs.

Public Interface:
    - OOXMLDocument: Wrapper around python-docx Document with OOXML access
    - get_xml_element: Extract XML elements by path
    - set_xml_property: Set XML properties directly
    - NAMESPACES: Common OOXML namespace dictionary

Example:
    >>> from docx_skill.ooxml import OOXMLDocument
    >>> 
    >>> # Create new document
    >>> doc = OOXMLDocument()
    >>> doc.add_paragraph("Hello World")
    >>> doc.save("output.docx")
    >>> 
    >>> # Load existing document
    >>> doc = OOXMLDocument.load("existing.docx")
    >>> 
    >>> # Access underlying python-docx Document
    >>> docx_obj = doc.document
    >>> 
    >>> # Direct XML access (advanced)
    >>> body = doc.get_body_element()
"""

from pathlib import Path
from typing import Optional, Any, List, Dict
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.xmlchemy import BaseOxmlElement
import xml.etree.ElementTree as ET
import os


# Path to modern template (Aptos font, Microsoft 365 defaults)
_MODERN_TEMPLATE_PATH = Path(__file__).parent / "templates" / "modern.docx"


# Common OOXML namespaces
NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'v': 'urn:schemas-microsoft-com:vml',
    'cp': 'http://schemas.openxmlformats.org/package/2006/metadata/core-properties',
    'dc': 'http://purl.org/dc/elements/1.1/',
    'dcterms': 'http://purl.org/dc/terms/',
}


class OOXMLDocument:
    """Wrapper around python-docx Document with OOXML access.
    
    Provides a unified interface for document operations with access to both
    high-level python-docx API and low-level OOXML manipulation.
    
    This class is the foundation for all document operations. Higher-level APIs
    (DocumentBuilder, AdvancedDocument) build on top of this.
    
    Attributes:
        document: Underlying python-docx Document object
        path: Path to document file (if loaded from/saved to file)
        
    Example:
        >>> # Create new document
        >>> doc = OOXMLDocument()
        >>> doc.add_paragraph("Content")
        >>> doc.save("output.docx")
        >>> 
        >>> # Load existing
        >>> doc = OOXMLDocument.load("existing.docx")
        >>> paragraphs = doc.get_paragraphs()
        >>> 
        >>> # Access python-docx directly
        >>> docx_obj = doc.document
        >>> for para in docx_obj.paragraphs:
        ...     print(para.text)
    """
    
    def __init__(self, document: Optional[Document] = None, use_modern_template: bool = True):
        """Initialize OOXMLDocument.
        
        Args:
            document: Existing python-docx Document (creates new if None)
            use_modern_template: Use modern Microsoft 365 template with Aptos font (default: True)
            
        Example:
            >>> # New document with modern template (Aptos font)
            >>> doc = OOXMLDocument()
            >>> 
            >>> # New document with legacy python-docx template (Calibri font)
            >>> doc = OOXMLDocument(use_modern_template=False)
            >>> 
            >>> # Wrap existing
            >>> from docx import Document
            >>> existing = Document("file.docx")
            >>> doc = OOXMLDocument(existing)
        """
        if document is not None:
            self.document = document
        elif use_modern_template and _MODERN_TEMPLATE_PATH.exists():
            # Use modern Microsoft 365 template with Aptos font
            self.document = Document(str(_MODERN_TEMPLATE_PATH))
        else:
            # Fall back to python-docx default template (Calibri font)
            self.document = Document()
        
        self.path: Optional[Path] = None
    
    @classmethod
    def load(cls, path: str | Path) -> 'OOXMLDocument':
        """Load document from file.
        
        Args:
            path: Path to DOCX file
            
        Returns:
            OOXMLDocument instance
            
        Raises:
            FileNotFoundError: If file doesn't exist
            
        Example:
            >>> doc = OOXMLDocument.load("document.docx")
            >>> print(f"Loaded {len(doc.get_paragraphs())} paragraphs")
        """
        path_obj = Path(path)
        if not path_obj.exists():
            raise FileNotFoundError(f"Document not found: {path}")
        
        document = Document(str(path_obj))
        instance = cls(document)
        instance.path = path_obj
        return instance
    
    def save(self, path: Optional[str | Path] = None) -> Path:
        """Save document to file.
        
        Args:
            path: Target path (uses self.path if None)
            
        Returns:
            Path where document was saved
            
        Raises:
            ValueError: If no path provided and no self.path set
            
        Example:
            >>> doc = OOXMLDocument()
            >>> doc.add_paragraph("Content")
            >>> doc.save("output.docx")
            >>> 
            >>> # Save to original location
            >>> doc = OOXMLDocument.load("file.docx")
            >>> doc.add_paragraph("More content")
            >>> doc.save()  # Updates file.docx
        """
        if path is None:
            if self.path is None:
                raise ValueError("No path specified and document not loaded from file")
            path = self.path
        
        path_obj = Path(path)
        path_obj.parent.mkdir(parents=True, exist_ok=True)
        
        self.document.save(str(path_obj))
        self.path = path_obj
        return path_obj
    
    # High-level content access
    
    def add_paragraph(self, text: str = "", style: Optional[str] = None) -> Any:
        """Add paragraph to document.
        
        Args:
            text: Paragraph text
            style: Style name (e.g., "Heading 1", "Normal")
            
        Returns:
            python-docx Paragraph object
            
        Example:
            >>> doc = OOXMLDocument()
            >>> doc.add_paragraph("Normal text")
            >>> doc.add_paragraph("Heading", style="Heading 1")
        """
        return self.document.add_paragraph(text, style)
    
    def add_heading(self, text: str, level: int = 1) -> Any:
        """Add heading to document.
        
        Args:
            text: Heading text
            level: Heading level (1-9, where 1 is largest)
            
        Returns:
            python-docx Paragraph object
            
        Example:
            >>> doc = OOXMLDocument()
            >>> doc.add_heading("Chapter 1", level=1)
            >>> doc.add_heading("Section 1.1", level=2)
        """
        return self.document.add_heading(text, level)
    
    def add_page_break(self) -> Any:
        """Add page break to document.
        
        Returns:
            python-docx Paragraph object containing page break
            
        Example:
            >>> doc = OOXMLDocument()
            >>> doc.add_paragraph("Page 1")
            >>> doc.add_page_break()
            >>> doc.add_paragraph("Page 2")
        """
        return self.document.add_page_break()
    
    def get_paragraphs(self) -> List[Any]:
        """Get all paragraphs in document.
        
        Returns:
            List of python-docx Paragraph objects
            
        Example:
            >>> doc = OOXMLDocument.load("document.docx")
            >>> for para in doc.get_paragraphs():
            ...     print(para.text)
        """
        return self.document.paragraphs
    
    def get_tables(self) -> List[Any]:
        """Get all tables in document.
        
        Returns:
            List of python-docx Table objects
            
        Example:
            >>> doc = OOXMLDocument.load("document.docx")
            >>> tables = doc.get_tables()
            >>> print(f"Found {len(tables)} tables")
        """
        return self.document.tables
    
    def get_sections(self) -> List[Any]:
        """Get all sections in document.
        
        Returns:
            List of python-docx Section objects
            
        Example:
            >>> doc = OOXMLDocument.load("document.docx")
            >>> for section in doc.get_sections():
            ...     print(f"Section: {section.orientation}")
        """
        return self.document.sections
    
    # Low-level OOXML access
    
    def get_body_element(self) -> BaseOxmlElement:
        """Get document body XML element.
        
        Returns:
            lxml Element representing document body
            
        Example:
            >>> doc = OOXMLDocument()
            >>> body = doc.get_body_element()
            >>> # Direct XML manipulation on body element
        """
        return self.document._element.body
    
    def get_document_element(self) -> BaseOxmlElement:
        """Get root document XML element.
        
        Returns:
            lxml Element representing document root
            
        Example:
            >>> doc = OOXMLDocument()
            >>> root = doc.get_document_element()
        """
        return self.document._element
    
    def find_elements(self, tag: str, namespace: Optional[str] = 'w') -> List[BaseOxmlElement]:
        """Find all elements matching tag in document.
        
        Args:
            tag: Tag name (without namespace prefix)
            namespace: Namespace key from NAMESPACES dict (default: 'w')
            
        Returns:
            List of matching XML elements
            
        Example:
            >>> doc = OOXMLDocument.load("document.docx")
            >>> # Find all paragraphs
            >>> paragraphs = doc.find_elements('p')
            >>> print(f"Found {len(paragraphs)} paragraph elements")
            >>> 
            >>> # Find all runs
            >>> runs = doc.find_elements('r')
        """
        if namespace:
            # qn() returns the qualified name with full namespace URI
            qualified_tag = qn(f'{namespace}:{tag}')
        else:
            qualified_tag = tag
        
        # Use findall with qualified name
        body = self.get_body_element()
        return body.findall(f'.//{qualified_tag}')
    
    def get_core_properties(self) -> Any:
        """Get document core properties (metadata).
        
        Returns:
            CoreProperties object with author, title, etc.
            
        Example:
            >>> doc = OOXMLDocument.load("document.docx")
            >>> props = doc.get_core_properties()
            >>> print(f"Author: {props.author}")
            >>> print(f"Title: {props.title}")
            >>> 
            >>> # Modify properties
            >>> props.author = "New Author"
            >>> props.title = "New Title"
            >>> doc.save()
        """
        return self.document.core_properties
    
    def get_styles(self) -> Any:
        """Get document styles.
        
        Returns:
            Styles object for accessing/modifying styles
            
        Example:
            >>> doc = OOXMLDocument.load("document.docx")
            >>> styles = doc.get_styles()
            >>> for style in styles:
            ...     print(style.name)
        """
        return self.document.styles
    
    def add_custom_xml(self, xml_string: str, namespace: str = 'w') -> BaseOxmlElement:
        """Add custom XML element to document body.
        
        Args:
            xml_string: XML string to parse and add
            namespace: Default namespace for element
            
        Returns:
            Added XML element
            
        Example:
            >>> doc = OOXMLDocument()
            >>> # Add custom paragraph with specific formatting
            >>> xml = '''
            ... <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            ...     <w:r><w:t>Custom XML content</w:t></w:r>
            ... </w:p>
            ... '''
            >>> doc.add_custom_xml(xml)
            >>> doc.save("custom.docx")
        """
        element = parse_xml(xml_string)
        self.get_body_element().append(element)
        return element
    
    # Utility methods
    
    def get_text(self) -> str:
        """Get all text content from document.
        
        Returns:
            Full document text with paragraphs joined by newlines
            
        Example:
            >>> doc = OOXMLDocument.load("document.docx")
            >>> text = doc.get_text()
            >>> print(f"Document has {len(text)} characters")
        """
        return '\n'.join(para.text for para in self.document.paragraphs)
    
    def get_word_count(self) -> int:
        """Get approximate word count.
        
        Returns:
            Number of words in document
            
        Example:
            >>> doc = OOXMLDocument.load("document.docx")
            >>> words = doc.get_word_count()
            >>> print(f"Document has {words:,} words")
        """
        text = self.get_text()
        return len(text.split())
    
    def clear_content(self) -> None:
        """Remove all content from document body.
        
        Warning:
            This removes ALL paragraphs, tables, and other elements.
            Use with caution.
            
        Example:
            >>> doc = OOXMLDocument.load("document.docx")
            >>> doc.clear_content()
            >>> doc.add_paragraph("Fresh start")
            >>> doc.save()
        """
        body = self.get_body_element()
        for element in list(body):
            body.remove(element)
    
    def __repr__(self) -> str:
        """String representation of document."""
        path_str = f"'{self.path}'" if self.path else "unsaved"
        para_count = len(self.get_paragraphs())
        word_count = self.get_word_count()
        return f"OOXMLDocument({path_str}, {para_count} paragraphs, {word_count} words)"


# Utility functions

def qualified_name(namespace: str, tag: str) -> str:
    """Get qualified XML name from namespace and tag.
    
    Args:
        namespace: Namespace key from NAMESPACES dict
        tag: Tag name without prefix
        
    Returns:
        Qualified name string
        
    Example:
        >>> qualified_name('w', 'p')
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'
    """
    return qn(f'{namespace}:{tag}')


def get_xml_element(element: BaseOxmlElement, path: str, namespaces: Optional[Dict[str, str]] = None) -> Optional[BaseOxmlElement]:
    """Find XML element by XPath.
    
    Args:
        element: Root element to search from
        path: XPath expression with namespace prefixes (e.g., './/w:p')
        namespaces: Namespace dict (uses NAMESPACES if None)
        
    Returns:
        First matching element or None
        
    Example:
        >>> doc = OOXMLDocument.load("document.docx")
        >>> body = doc.get_body_element()
        >>> # Convert namespace prefix to qualified name
        >>> first_para = get_xml_element(body, './/w:p', NAMESPACES)
    """
    # Convert XPath with namespace prefixes to qualified names
    # For simple paths like './/w:p', extract namespace prefix and tag
    if '/' in path:
        parts = path.split('/')
        for i, part in enumerate(parts):
            if ':' in part and part != '.':
                prefix, tag = part.split(':', 1)
                if prefix in NAMESPACES:
                    # Replace with qualified name
                    parts[i] = qn(f'{prefix}:{tag}')
        path = '/'.join(parts)
    
    results = element.findall(path)
    return results[0] if results else None


def get_xml_elements(element: BaseOxmlElement, path: str, namespaces: Optional[Dict[str, str]] = None) -> List[BaseOxmlElement]:
    """Find all XML elements by XPath.
    
    Args:
        element: Root element to search from
        path: XPath expression with namespace prefixes (e.g., './/w:p')
        namespaces: Namespace dict (uses NAMESPACES if None)
        
    Returns:
        List of matching elements
        
    Example:
        >>> doc = OOXMLDocument.load("document.docx")
        >>> body = doc.get_body_element()
        >>> all_paras = get_xml_elements(body, './/w:p', NAMESPACES)
        >>> print(f"Found {len(all_paras)} paragraphs")
    """
    # Convert XPath with namespace prefixes to qualified names
    if '/' in path:
        parts = path.split('/')
        for i, part in enumerate(parts):
            if ':' in part and part != '.':
                prefix, tag = part.split(':', 1)
                if prefix in NAMESPACES:
                    # Replace with qualified name
                    parts[i] = qn(f'{prefix}:{tag}')
        path = '/'.join(parts)
    
    return element.findall(path)


def set_xml_property(element: BaseOxmlElement, property_path: str, value: str, namespaces: Optional[Dict[str, str]] = None) -> bool:
    """Set XML property value via XPath.
    
    Args:
        element: Root element
        property_path: XPath to property
        value: Value to set
        namespaces: Namespace dict (uses NAMESPACES if None)
        
    Returns:
        True if property was set, False if not found
        
    Example:
        >>> doc = OOXMLDocument()
        >>> para = doc.add_paragraph("Text")
        >>> # Set paragraph alignment to center
        >>> set_xml_property(para._element, './/w:pPr/w:jc', 'center')
    """
    ns = namespaces if namespaces is not None else NAMESPACES
    target = element.xpath(property_path, namespaces=ns)
    
    if target:
        if hasattr(target[0], 'set'):
            target[0].set('val', value)
        else:
            target[0].text = value
        return True
    
    return False
