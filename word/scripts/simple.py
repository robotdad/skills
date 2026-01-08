"""
Module: Simple API - DocumentBuilder

High-level builder pattern for 80% of document creation use cases.
Provides a fluent interface with method chaining for quick document creation.

Public Interface:
    - DocumentBuilder: High-level document builder with safe defaults

Example:
    >>> from docx_skill.simple import DocumentBuilder
    >>> 
    >>> # Create document with fluent interface
    >>> builder = DocumentBuilder()
    >>> builder.add_heading("Report Title", level=1)
    >>> builder.add_paragraph("Introduction text", bold=True)
    >>> builder.add_list(["Item 1", "Item 2", "Item 3"])
    >>> builder.save("report.docx")
    >>> 
    >>> # Method chaining
    >>> (DocumentBuilder()
    ...     .add_heading("Quick Report")
    ...     .add_paragraph("First paragraph")
    ...     .add_paragraph("Second paragraph")
    ...     .save("quick.docx"))
"""

from pathlib import Path
from typing import Optional, List, Any
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from .ooxml import OOXMLDocument
    from .safety import SafeFileOperations
    from .validation import validate_docx
except ImportError:
    # Fallback for direct imports (e.g., in tests)
    from ooxml import OOXMLDocument
    from safety import SafeFileOperations
    from validation import validate_docx


class DocumentBuilder:
    """High-level document builder with safe defaults.
    
    Provides a fluent interface for common document creation tasks. Methods
    return self to enable method chaining. Uses modern Microsoft 365 template
    with Aptos font by default.
    
    This is the recommended API for most document creation tasks. For advanced
    control over styles, sections, and formatting, use AdvancedDocument instead.
    
    Attributes:
        _doc: Underlying OOXMLDocument instance
        _safe_ops: SafeFileOperations for safe saving
        
    Example:
        >>> # Simple document
        >>> doc = DocumentBuilder()
        >>> doc.add_heading("My Document")
        >>> doc.add_paragraph("Hello World")
        >>> doc.save("output.docx")
        >>> 
        >>> # Method chaining
        >>> (DocumentBuilder()
        ...     .add_heading("Report", level=1)
        ...     .add_paragraph("Introduction", bold=True)
        ...     .add_list(["Point 1", "Point 2"])
        ...     .add_page_break()
        ...     .add_heading("Conclusion", level=1)
        ...     .add_paragraph("Summary text")
        ...     .save("report.docx"))
    """
    
    def __init__(self, template: Optional[str] = None):
        """Create new document builder.
        
        Args:
            template: Optional template name or path. If None, uses modern
                     Microsoft 365 template with Aptos font. Special values:
                     - None or "modern": Modern Microsoft 365 template (default)
                     - "legacy": Python-docx default template (Calibri)
                     - Any path: Custom template file
        
        Example:
            >>> # Default modern template
            >>> builder = DocumentBuilder()
            >>> 
            >>> # Legacy template
            >>> builder = DocumentBuilder(template="legacy")
            >>> 
            >>> # Custom template
            >>> builder = DocumentBuilder(template="path/to/template.docx")
        """
        if template is None or template == "modern":
            # Use modern template (default)
            self._doc = OOXMLDocument(use_modern_template=True)
        elif template == "legacy":
            # Use legacy python-docx template
            self._doc = OOXMLDocument(use_modern_template=False)
        else:
            # Load custom template
            self._doc = OOXMLDocument.load(template)
        
        self._safe_ops = SafeFileOperations(default_allow_overwrite=False)
        
        # Cache for list numId lookups to avoid repeated XML parsing
        self._list_numids_cache: Optional[tuple] = None
    
    def add_heading(self, text: str, level: int = 1) -> 'DocumentBuilder':
        """Add heading with specified level.
        
        Args:
            text: Heading text
            level: Heading level (1-9, where 1 is largest)
        
        Returns:
            Self for method chaining
            
        Example:
            >>> builder = DocumentBuilder()
            >>> builder.add_heading("Chapter 1", level=1)
            >>> builder.add_heading("Section 1.1", level=2)
            >>> builder.add_heading("Subsection 1.1.1", level=3)
        """
        self._doc.add_heading(text, level=level)
        return self
    
    def add_paragraph(self, text: str, 
                     bold: bool = False,
                     italic: bool = False,
                     font_size: Optional[int] = None) -> 'DocumentBuilder':
        """Add paragraph with basic formatting.
        
        Args:
            text: Paragraph text
            bold: Make text bold
            italic: Make text italic
            font_size: Font size in points (e.g., 12)
        
        Returns:
            Self for method chaining
            
        Example:
            >>> builder = DocumentBuilder()
            >>> builder.add_paragraph("Normal text")
            >>> builder.add_paragraph("Bold text", bold=True)
            >>> builder.add_paragraph("Italic text", italic=True)
            >>> builder.add_paragraph("Large text", font_size=18)
            >>> builder.add_paragraph("Bold and italic", bold=True, italic=True)
        """
        para = self._doc.add_paragraph(text)
        
        # Apply formatting to all runs in paragraph
        for run in para.runs:
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if font_size:
                run.font.size = Pt(font_size)
        
        return self
    
    def _get_list_numids(self) -> tuple:
        """
        Find numId values for bullet and decimal lists in document.
        
        This dynamically looks up the correct numId values by examining the
        numbering.xml part. This is necessary because python-docx reorganizes
        numbering.xml when saving, causing hardcoded numId values to become
        invalid.
        
        Returns:
            tuple: (bullet_numId, decimal_numId) or (None, None) if not found
            
        Note:
            Results are cached after first call to avoid repeated XML parsing.
        """
        # Return cached values if available
        if self._list_numids_cache is not None:
            return self._list_numids_cache
        
        from docx.oxml.ns import qn
        
        bullet_numid = None
        decimal_numid = None
        
        try:
            # Access the numbering part
            numbering_part = self._doc.document.part.numbering_part
            if numbering_part is None:
                # No numbering part exists - document has no list definitions
                self._list_numids_cache = (None, None)
                return self._list_numids_cache
            
            # Get the numbering element (root of numbering.xml)
            numbering = numbering_part.element
            
            # Build a map of abstractNumId -> format
            abstract_formats = {}
            for abstract_num in numbering.findall(qn('w:abstractNum')):
                abstract_id = abstract_num.get(qn('w:abstractNumId'))
                
                # Get the format from level 0
                lvl = abstract_num.find(f'.//{qn("w:lvl")}[@{qn("w:ilvl")}="0"]')
                if lvl is not None:
                    num_fmt = lvl.find(qn('w:numFmt'))
                    if num_fmt is not None:
                        fmt_val = num_fmt.get(qn('w:val'))
                        abstract_formats[abstract_id] = fmt_val
            
            # Find the first bullet and decimal numId values
            for num in numbering.findall(qn('w:num')):
                num_id = num.get(qn('w:numId'))
                
                # Get the abstractNumId this num references
                abstract_num_id_elem = num.find(qn('w:abstractNumId'))
                if abstract_num_id_elem is not None:
                    abstract_id = abstract_num_id_elem.get(qn('w:val'))
                    fmt = abstract_formats.get(abstract_id)
                    
                    # Check if this is a bullet or decimal format
                    if fmt == 'bullet' and bullet_numid is None:
                        bullet_numid = int(num_id)
                    elif fmt == 'decimal' and decimal_numid is None:
                        decimal_numid = int(num_id)
                    
                    # Stop if we found both
                    if bullet_numid is not None and decimal_numid is not None:
                        break
        
        except Exception as e:
            # If anything goes wrong, return None values
            # This allows fallback to character-based lists
            print(f"Warning: Could not determine list numIds: {e}")
            self._list_numids_cache = (None, None)
            return self._list_numids_cache
        
        # Cache the results
        self._list_numids_cache = (bullet_numid, decimal_numid)
        return self._list_numids_cache
    
    def add_list(self, items: List[str], numbered: bool = False) -> 'DocumentBuilder':
        """Add bulleted or numbered list.
        
        Creates real Word lists with proper numPr elements for correct
        markdown conversion. Dynamically looks up numId values from the
        document's numbering.xml to ensure correct list formatting even
        after python-docx reorganizes the numbering definitions.
        
        Falls back to character-based lists if numbering definitions
        are not found.
        
        Args:
            items: List of item texts
            numbered: If True, create numbered list; otherwise bulleted
        
        Returns:
            Self for method chaining
            
        Example:
            >>> builder = DocumentBuilder()
            >>> # Bulleted list
            >>> builder.add_list([
            ...     "First item",
            ...     "Second item",
            ...     "Third item"
            ... ])
            >>> 
            >>> # Numbered list
            >>> builder.add_list([
            ...     "Step 1",
            ...     "Step 2",
            ...     "Step 3"
            ... ], numbered=True)
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        # Dynamically look up numId values from numbering.xml
        bullet_numid, decimal_numid = self._get_list_numids()
        
        # Determine which numId to use
        if numbered:
            num_id = decimal_numid
            fallback_char = '1.'
        else:
            num_id = bullet_numid
            fallback_char = '•'
        
        # If no numbering definitions exist, fall back to character-based lists
        if num_id is None:
            for item in items:
                self._doc.add_paragraph(f"{fallback_char} {item}")
            return self
        
        # Create real Word list items with numPr elements
        for item in items:
            # Add paragraph
            para = self._doc.add_paragraph(item)
            
            # Add numPr element to make it a real Word list item
            pPr = para._element.get_or_add_pPr()
            
            numPr = OxmlElement('w:numPr')
            
            # ilvl = indent level (0 for first level)
            ilvl = OxmlElement('w:ilvl')
            ilvl.set(qn('w:val'), '0')
            numPr.append(ilvl)
            
            # numId = numbering definition reference
            numId_elem = OxmlElement('w:numId')
            numId_elem.set(qn('w:val'), str(num_id))
            numPr.append(numId_elem)
            
            pPr.append(numPr)
        
        return self
    
    def add_table(self, data: List[List[str]], 
                 headers: Optional[List[str]] = None) -> 'DocumentBuilder':
        """Add simple table from 2D array.
        
        Args:
            data: 2D list of cell values (rows x columns)
            headers: Optional header row. If provided, will be formatted
                    as table header (bold, shaded background)
        
        Returns:
            Self for method chaining
            
        Example:
            >>> builder = DocumentBuilder()
            >>> 
            >>> # Simple table without headers
            >>> builder.add_table([
            ...     ["A1", "B1", "C1"],
            ...     ["A2", "B2", "C2"],
            ...     ["A3", "B3", "C3"]
            ... ])
            >>> 
            >>> # Table with headers
            >>> builder.add_table(
            ...     data=[
            ...         ["John", "Doe", "30"],
            ...         ["Jane", "Smith", "25"]
            ...     ],
            ...     headers=["First Name", "Last Name", "Age"]
            ... )
        """
        if not data:
            return self
        
        # Determine table dimensions
        num_cols = len(data[0])
        num_rows = len(data) + (1 if headers else 0)
        
        # Create table
        table = self._doc.document.add_table(rows=num_rows, cols=num_cols)
        
        # Try to apply modern table style if it exists
        try:
            table.style = 'Light Grid Accent 1'
        except KeyError:
            # Style doesn't exist, use default table style
            pass
        
        # Add headers if provided
        row_offset = 0
        if headers:
            header_row = table.rows[0]
            for col_idx, header_text in enumerate(headers):
                cell = header_row.cells[col_idx]
                cell.text = header_text
                # Make header bold
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            row_offset = 1
        
        # Add data rows
        for row_idx, row_data in enumerate(data):
            table_row = table.rows[row_idx + row_offset]
            for col_idx, cell_value in enumerate(row_data):
                table_row.cells[col_idx].text = str(cell_value)
        
        return self
    
    def add_image(self, path: str, 
                 width_inches: Optional[float] = None) -> 'DocumentBuilder':
        """Add image with optional sizing.
        
        Args:
            path: Path to image file (PNG, JPG, etc.)
            width_inches: Image width in inches. If None, uses original size.
                         Height will be scaled proportionally.
        
        Returns:
            Self for method chaining
            
        Raises:
            FileNotFoundError: If image file doesn't exist
            
        Example:
            >>> builder = DocumentBuilder()
            >>> # Add image at original size
            >>> builder.add_image("photo.jpg")
            >>> 
            >>> # Add image with specific width
            >>> builder.add_image("chart.png", width_inches=5)
            >>> 
            >>> # Add multiple images
            >>> builder.add_heading("Images")
            >>> builder.add_image("image1.jpg", width_inches=4)
            >>> builder.add_paragraph("Figure 1: Description")
            >>> builder.add_image("image2.jpg", width_inches=4)
        """
        image_path = Path(path)
        if not image_path.exists():
            raise FileNotFoundError(f"Image not found: {path}")
        
        # Add image to document
        if width_inches:
            self._doc.document.add_picture(str(image_path), width=Inches(width_inches))
        else:
            self._doc.document.add_picture(str(image_path))
        
        return self
    
    def add_page_break(self) -> 'DocumentBuilder':
        """Add page break.
        
        Returns:
            Self for method chaining
            
        Example:
            >>> builder = DocumentBuilder()
            >>> builder.add_paragraph("Page 1 content")
            >>> builder.add_page_break()
            >>> builder.add_paragraph("Page 2 content")
            >>> builder.add_page_break()
            >>> builder.add_paragraph("Page 3 content")
        """
        self._doc.add_page_break()
        return self
    
    def save(self, output_path: str, overwrite: bool = False) -> str:
        """Save document to path.
        
        Args:
            output_path: Where to save document
            overwrite: If False (default), raises error if file exists.
                      If True, overwrites existing file with backup.
        
        Returns:
            Actual path where document was saved
            
        Raises:
            FileExistsError: If file exists and overwrite=False
            
        Example:
            >>> builder = DocumentBuilder()
            >>> builder.add_heading("Document")
            >>> builder.add_paragraph("Content")
            >>> 
            >>> # Save with overwrite protection (default)
            >>> path = builder.save("output.docx")
            >>> 
            >>> # Save with overwrite allowed
            >>> path = builder.save("output.docx", overwrite=True)
        """
        # Save document to temp location first
        temp_path = Path(output_path).parent / f"~temp_{Path(output_path).name}"
        self._doc.save(temp_path)
        
        # Validate the saved document
        try:
            validation = validate_docx(str(temp_path))
            if not validation.is_valid:
                # Log warnings but don't fail
                for issue in validation.issues:
                    print(f"Warning: {issue.message}")
        except Exception as e:
            # Don't fail save on validation errors
            print(f"Warning: Validation failed: {e}")
        
        # Use SafeFileOperations to move to final location
        try:
            data = temp_path.read_bytes()
            self._safe_ops.write_file(
                data=data,
                target_path=output_path,
                allow_overwrite=overwrite,
                backup=overwrite
            )
        finally:
            # Clean up temp file
            if temp_path.exists():
                temp_path.unlink()
        
        return str(output_path)
    
    def get_document(self) -> OOXMLDocument:
        """Get underlying OOXMLDocument for advanced operations.
        
        Returns:
            OOXMLDocument instance
            
        Example:
            >>> builder = DocumentBuilder()
            >>> builder.add_heading("Title")
            >>> 
            >>> # Access underlying document for advanced operations
            >>> doc = builder.get_document()
            >>> props = doc.get_core_properties()
            >>> props.author = "John Doe"
            >>> props.title = "My Document"
        """
        return self._doc
    
    def __repr__(self) -> str:
        """String representation of builder."""
        para_count = len(self._doc.get_paragraphs())
        table_count = len(self._doc.get_tables())
        return f"DocumentBuilder({para_count} paragraphs, {table_count} tables)"
