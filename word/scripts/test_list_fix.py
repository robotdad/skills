#!/usr/bin/env python3
"""
Test script to verify dynamic numId lookup fix for lists.

This tests that:
1. Both bulleted and numbered lists are created correctly
2. After saving, lists still use the correct format
3. Lists convert to proper markdown format
"""

import sys
from pathlib import Path
from simple import DocumentBuilder

def test_lists():
    """Test bullet and numbered list creation."""
    print("=" * 60)
    print("Testing List Fix: Dynamic numId Lookup")
    print("=" * 60)
    
    # Create test document
    print("\n1. Creating document with bullets and numbered lists...")
    builder = DocumentBuilder()
    
    # Add heading
    builder.add_heading("List Testing Document", level=1)
    
    # Add bulleted list
    builder.add_heading("Bulleted List", level=2)
    builder.add_list([
        "First bullet item",
        "Second bullet item",
        "Third bullet item"
    ], numbered=False)
    
    # Add numbered list
    builder.add_heading("Numbered List", level=2)
    builder.add_list([
        "First numbered item",
        "Second numbered item", 
        "Third numbered item"
    ], numbered=True)
    
    # Save document
    output_path = "/Users/robotdad/guide/test_lists.docx"
    print(f"\n2. Saving document to: {output_path}")
    saved_path = builder.save(output_path, overwrite=True)
    print(f"   ✓ Saved successfully to: {saved_path}")
    
    # Verify the document
    print("\n3. Verifying list numId values...")
    
    # Re-load the document to see how python-docx reorganized it
    from docx import Document
    doc = Document(saved_path)
    
    # Check numbering part exists
    if doc.part.numbering_part is None:
        print("   ✗ ERROR: No numbering part found!")
        return False
    
    print("   ✓ Numbering part exists")
    
    # Analyze the numbering definitions
    from docx.oxml.ns import qn
    numbering = doc.part.numbering_part.element
    
    # Build abstract format map
    abstract_formats = {}
    for abstract_num in numbering.findall(qn('w:abstractNum')):
        abstract_id = abstract_num.get(qn('w:abstractNumId'))
        lvl = abstract_num.find(f'.//{qn("w:lvl")}[@{qn("w:ilvl")}="0"]')
        if lvl is not None:
            num_fmt = lvl.find(qn('w:numFmt'))
            if num_fmt is not None:
                fmt_val = num_fmt.get(qn('w:val'))
                abstract_formats[abstract_id] = fmt_val
    
    print(f"   ✓ Found {len(abstract_formats)} abstract numbering definitions")
    
    # Map numId to format
    numid_to_format = {}
    for num in numbering.findall(qn('w:num')):
        num_id = num.get(qn('w:numId'))
        abstract_num_id_elem = num.find(qn('w:abstractNumId'))
        if abstract_num_id_elem is not None:
            abstract_id = abstract_num_id_elem.get(qn('w:val'))
            fmt = abstract_formats.get(abstract_id)
            if fmt:
                numid_to_format[num_id] = fmt
                print(f"     numId={num_id} → format={fmt}")
    
    # Check paragraphs for list items
    print("\n4. Checking list items in document...")
    bullet_items = []
    decimal_items = []
    
    for para in doc.paragraphs:
        # Check if paragraph has numPr
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                numId_elem = numPr.find(qn('w:numId'))
                if numId_elem is not None:
                    num_id = numId_elem.get(qn('w:val'))
                    fmt = numid_to_format.get(num_id)
                    text = para.text
                    
                    if fmt == 'bullet':
                        bullet_items.append(text)
                        print(f"   ✓ Bullet item (numId={num_id}): {text}")
                    elif fmt == 'decimal':
                        decimal_items.append(text)
                        print(f"   ✓ Decimal item (numId={num_id}): {text}")
                    else:
                        print(f"   ✗ Unknown format '{fmt}' for: {text}")
    
    # Verify counts
    print("\n5. Verification Results:")
    success = True
    
    if len(bullet_items) == 3:
        print(f"   ✓ Found {len(bullet_items)} bullet items (expected 3)")
    else:
        print(f"   ✗ Found {len(bullet_items)} bullet items (expected 3)")
        success = False
    
    if len(decimal_items) == 3:
        print(f"   ✓ Found {len(decimal_items)} decimal items (expected 3)")
    else:
        print(f"   ✗ Found {len(decimal_items)} decimal items (expected 3)")
        success = False
    
    # Test markdown conversion
    print("\n6. Testing markdown conversion...")
    expected_bullets = [
        "* First bullet item",
        "* Second bullet item",
        "* Third bullet item"
    ]
    expected_numbers = [
        "1. First numbered item",
        "2. Second numbered item",
        "3. Third numbered item"
    ]
    
    # Import markdown converter
    try:
        from md_converter import DocxToMarkdown
        
        converter = DocxToMarkdown()
        markdown = converter.convert(saved_path)
        
        # Check for bullet markers
        bullet_match = all(item in markdown for item in expected_bullets)
        if bullet_match:
            print("   ✓ Bullets convert to markdown correctly")
        else:
            print("   ✗ Bullets do not convert correctly")
            print(f"     Looking for: {expected_bullets}")
            success = False
        
        # Check for number markers
        number_match = all(item in markdown for item in expected_numbers)
        if number_match:
            print("   ✓ Numbers convert to markdown correctly")
        else:
            print("   ✗ Numbers do not convert correctly")
            print(f"     Looking for: {expected_numbers}")
            success = False
        
        # Print markdown for inspection
        print("\n7. Generated Markdown:")
        print("-" * 60)
        print(markdown)
        print("-" * 60)
        
    except ImportError:
        print("   ⚠ Skipping markdown test (md_converter not available)")
    
    print("\n" + "=" * 60)
    if success:
        print("✓ ALL TESTS PASSED")
        print("=" * 60)
        print(f"\nDocument saved to: {saved_path}")
        print("Open in Word to verify lists appear correctly.")
        return True
    else:
        print("✗ SOME TESTS FAILED")
        print("=" * 60)
        return False

if __name__ == "__main__":
    success = test_lists()
    sys.exit(0 if success else 1)
