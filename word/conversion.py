"""
Module: Format Conversion

Provides utilities for converting DOCX documents to Markdown and extracting plain text.
Uses markitdown for DOCX → Markdown conversion (optimized for LLM consumption).

markitdown is specifically designed for AI/RAG workflows and requires no system 
dependencies - just pip install markitdown.

Public Interface:
    - docx_to_markdown: Convert DOCX to Markdown using markitdown
    - extract_text: Extract plain text from DOCX (no formatting)
    - is_markitdown_available: Check if markitdown is installed

Example:
    >>> from docx_skill.conversion import docx_to_markdown, extract_text
    >>> 
    >>> # Convert DOCX to Markdown
    >>> markdown_content = docx_to_markdown("report.docx", "report.md")
    >>> print(f"Converted to {len(markdown_content)} characters")
    >>> 
    >>> # Extract plain text for analysis
    >>> text = extract_text("document.docx")
    >>> word_count = len(text.split())
    >>> print(f"Document has {word_count} words")
"""

import os
import shutil
from pathlib import Path
from typing import Optional
from docx import Document

try:
    from .safety import SafeFileOperations
except ImportError:
    from safety import SafeFileOperations


def is_markitdown_available() -> bool:
    """Check if markitdown library is installed.
    
    markitdown is a lightweight library that requires no system dependencies.
    It's specifically designed for converting documents to Markdown for LLM consumption.
    
    Returns:
        True if markitdown is available, False otherwise
        
    Example:
        >>> if is_markitdown_available():
        ...     markdown = docx_to_markdown("doc.docx")
        ... else:
        ...     print("Install markitdown: pip install 'markitdown[all]'")
    """
    try:
        import markitdown
        return True
    except ImportError:
        return False


def docx_to_markdown(
    docx_path: str | Path,
    output_path: Optional[str | Path] = None
) -> str:
    """Convert DOCX to Markdown using markitdown (optimized for LLMs).
    
    Uses Microsoft's markitdown library which is specifically designed for
    AI/RAG workflows, preserving document structure perfectly for LLM consumption.
    No system dependencies required - just pip install markitdown.
    
    Args:
        docx_path: Path to DOCX file to convert
        output_path: Optional path to save markdown file. If None, only returns string.
        
    Returns:
        Markdown content as string
        
    Raises:
        FileNotFoundError: If docx_path doesn't exist
        ImportError: If markitdown not installed
        ValueError: If conversion fails
        
    Example:
        >>> # Convert and save to file
        >>> markdown = docx_to_markdown("report.docx", "report.md")
        >>> print(f"Converted to: {len(markdown)} characters")
        >>> 
        >>> # Convert to string only (no file)
        >>> markdown_text = docx_to_markdown("notes.docx")
        >>> print(markdown_text[:100])  # Preview first 100 chars
        >>> 
        >>> # Use in LLM workflow
        >>> markdown = docx_to_markdown("contract.docx")
        >>> # Now pass markdown to your LLM for analysis
    """
    # Check if markitdown is available
    if not is_markitdown_available():
        raise ImportError(
            "markitdown is not installed.\n"
            "Install with: pip install markitdown\n"
            "Or with all features: pip install 'markitdown[all]'"
        )
    
    from markitdown import MarkItDown
    
    # Validate input file
    docx_file = Path(docx_path)
    if not docx_file.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_file}")
    
    if not docx_file.is_file():
        raise ValueError(f"Path is not a file: {docx_file}")
    
    try:
        # Convert using markitdown
        md = MarkItDown()
        result = md.convert(str(docx_file))
        markdown_content = result.text_content
        
        # Save to file if output_path provided
        if output_path:
            output_file = Path(output_path)
            safe_ops = SafeFileOperations()
            safe_ops.write_file(
                markdown_content.encode('utf-8'),
                output_file,
                allow_overwrite=True,
                backup=True
            )
        
        return markdown_content
        
    except Exception as e:
        raise ValueError(f"Failed to convert DOCX to Markdown: {e}") from e


def extract_text(docx_path: str | Path) -> str:
    """Extract plain text from DOCX (no formatting).
    
    Extracts all text content from a DOCX document without any formatting,
    styles, or structure. Useful for text analysis, search, or processing.
    Includes text from both paragraphs and tables.
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        Plain text content (paragraphs separated by newlines)
        
    Raises:
        FileNotFoundError: If docx_path doesn't exist
        ValueError: If file is not a valid DOCX document
        
    Example:
        >>> # Extract text for analysis
        >>> text = extract_text("document.docx")
        >>> word_count = len(text.split())
        >>> print(f"Document has {word_count} words")
        >>> 
        >>> # Search for keywords
        >>> text = extract_text("report.docx")
        >>> if "conclusion" in text.lower():
        ...     print("Document contains conclusion section")
        >>> 
        >>> # Extract for LLM processing
        >>> text = extract_text("contract.docx")
        >>> # Pass text to LLM for analysis
    """
    # Validate input file
    docx_file = Path(docx_path)
    if not docx_file.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_file}")
    
    if not docx_file.is_file():
        raise ValueError(f"Path is not a file: {docx_file}")
    
    try:
        # Load document and extract text
        doc = Document(docx_file)
        
        # Extract text from all paragraphs
        paragraphs = [para.text for para in doc.paragraphs]
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        
        # Join with newlines
        text = "\n".join(paragraphs)
        
        return text
        
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {e}") from e


# Public exports
__all__ = [
    'docx_to_markdown',
    'extract_text',
    'is_markitdown_available',
]
